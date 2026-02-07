#!/usr/bin/env python3
"""
AI Content Humanizer
Detects AI-generated content and humanizes it using paraphrasing models
"""

import PyPDF2
import torch
from transformers import (
    AutoTokenizer, 
    AutoModelForSequenceClassification,
    AutoModelForSeq2SeqLM,
    pipeline
)
import numpy as np
from typing import List, Tuple, Dict
import argparse
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

class AIContentHumanizer:
    def __init__(self, 
                 detector_model: str = "roberta-base-openai-detector",
                 paraphrase_model: str = "humarin/chatgpt_paraphraser_on_T5_base"):
        """
        Initialize the AI content humanizer
        
        Args:
            detector_model: Model for AI detection
            paraphrase_model: Model for paraphrasing/humanizing
        """
        print("Loading AI detection model...")
        try:
            self.detector_tokenizer = AutoTokenizer.from_pretrained(detector_model)
            self.detector_model = AutoModelForSequenceClassification.from_pretrained(detector_model)
            self.detector_model.eval()
        except Exception as e:
            print(f"Error loading detector: {e}")
            print("Trying alternative model...")
            detector_model = "Hello-SimpleAI/chatgpt-detector-roberta"
            self.detector_tokenizer = AutoTokenizer.from_pretrained(detector_model)
            self.detector_model = AutoModelForSequenceClassification.from_pretrained(detector_model)
            self.detector_model.eval()
        
        print("Loading paraphrasing model...")
        try:
            self.paraphrase_tokenizer = AutoTokenizer.from_pretrained(paraphrase_model)
            self.paraphrase_model = AutoModelForSeq2SeqLM.from_pretrained(paraphrase_model)
            self.paraphrase_model.eval()
        except Exception as e:
            print(f"Error loading paraphraser: {e}")
            print("Trying alternative model...")
            paraphrase_model = "Vamsi/T5_Paraphrase_Paws"
            self.paraphrase_tokenizer = AutoTokenizer.from_pretrained(paraphrase_model)
            self.paraphrase_model = AutoModelForSeq2SeqLM.from_pretrained(paraphrase_model)
            self.paraphrase_model.eval()
        
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        self.detector_model.to(self.device)
        self.paraphrase_model.to(self.device)
        print(f"Models loaded on {self.device}")
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text content from PDF file"""
        print(f"Extracting text from {pdf_path}...")
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text += page.extract_text() + "\n"
                    if (page_num + 1) % 5 == 0:
                        print(f"Processed {page_num + 1} pages...")
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
        
        print(f"Extracted {len(text.split())} words")
        return text
    
    def split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences while preserving structure"""
        # Clean text
        text = re.sub(r'\s+', ' ', text)
        # Split by sentence endings
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def detect_ai_probability(self, text: str) -> float:
        """Detect AI probability for a text segment"""
        inputs = self.detector_tokenizer(
            text,
            return_tensors="pt",
            truncation=True,
            max_length=512,
            padding=True
        ).to(self.device)
        
        with torch.no_grad():
            outputs = self.detector_model(**inputs)
            predictions = torch.softmax(outputs.logits, dim=-1)
            # Assuming label 1 is "AI-generated"
            ai_probability = predictions[0][1].item()
        
        return ai_probability
    
    def paraphrase_text(self, text: str, num_variants: int = 3) -> List[str]:
        """
        Paraphrase text to make it more human-like
        
        Args:
            text: Text to paraphrase
            num_variants: Number of paraphrase variants to generate
        
        Returns:
            List of paraphrased versions
        """
        # Prepare input
        input_text = f"paraphrase: {text}"
        
        inputs = self.paraphrase_tokenizer(
            input_text,
            return_tensors="pt",
            truncation=True,
            max_length=512,
            padding=True
        ).to(self.device)
        
        # Generate paraphrases
        with torch.no_grad():
            outputs = self.paraphrase_model.generate(
                **inputs,
                max_length=512,
                num_return_sequences=num_variants,
                num_beams=num_variants,
                temperature=0.8,
                do_sample=True,
                top_k=50,
                top_p=0.95,
                early_stopping=True
            )
        
        paraphrases = []
        for output in outputs:
            paraphrase = self.paraphrase_tokenizer.decode(output, skip_special_tokens=True)
            paraphrases.append(paraphrase.strip())
        
        return paraphrases
    
    def humanize_sentence(self, sentence: str, threshold: float = 0.7) -> Dict:
        """
        Humanize a sentence if it's detected as AI-generated
        
        Returns:
            Dict with original, ai_probability, humanized versions
        """
        # Detect AI probability
        ai_prob = self.detect_ai_probability(sentence)
        
        result = {
            'original': sentence,
            'ai_probability': ai_prob,
            'needs_humanization': ai_prob > threshold,
            'humanized_versions': []
        }
        
        # If AI probability is high, paraphrase
        if ai_prob > threshold:
            paraphrases = self.paraphrase_text(sentence, num_variants=3)
            
            # Check AI probability of paraphrases and select best
            scored_paraphrases = []
            for para in paraphrases:
                para_ai_prob = self.detect_ai_probability(para)
                scored_paraphrases.append({
                    'text': para,
                    'ai_probability': para_ai_prob,
                    'improvement': ai_prob - para_ai_prob
                })
            
            # Sort by lowest AI probability
            scored_paraphrases.sort(key=lambda x: x['ai_probability'])
            result['humanized_versions'] = scored_paraphrases
        
        return result
    
    def process_document(self, text: str, threshold: float = 0.7) -> Dict:
        """
        Process entire document and humanize AI-detected content
        
        Args:
            text: Document text
            threshold: AI probability threshold (0-1)
        
        Returns:
            Dict with processing results
        """
        sentences = self.split_into_sentences(text)
        print(f"\nProcessing {len(sentences)} sentences...")
        print(f"AI detection threshold: {threshold*100:.0f}%")
        
        results = {
            'total_sentences': len(sentences),
            'flagged_sentences': 0,
            'sentences_data': [],
            'statistics': {}
        }
        
        for i, sentence in enumerate(sentences):
            if len(sentence.split()) < 5:  # Skip very short sentences
                results['sentences_data'].append({
                    'original': sentence,
                    'ai_probability': 0.0,
                    'needs_humanization': False,
                    'humanized_versions': []
                })
                continue
            
            # Process sentence
            result = self.humanize_sentence(sentence, threshold)
            results['sentences_data'].append(result)
            
            if result['needs_humanization']:
                results['flagged_sentences'] += 1
            
            # Progress update
            if (i + 1) % 10 == 0:
                print(f"Processed {i + 1}/{len(sentences)} sentences... "
                      f"(Flagged: {results['flagged_sentences']})")
        
        # Calculate statistics
        ai_probs = [s['ai_probability'] for s in results['sentences_data']]
        results['statistics'] = {
            'average_ai_probability': np.mean(ai_probs),
            'flagged_percentage': (results['flagged_sentences'] / len(sentences)) * 100,
            'max_ai_probability': max(ai_probs),
            'min_ai_probability': min(ai_probs)
        }
        
        return results
    
    def create_comparison_docx(self, results: Dict, output_file: str):
        """Create a Word document showing original vs humanized text"""
        print(f"\nCreating comparison document: {output_file}")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('AI Content Humanization Report', 0)
        
        # Statistics
        doc.add_heading('Summary Statistics', 1)
        stats = results['statistics']
        doc.add_paragraph(f"Total sentences analyzed: {results['total_sentences']}")
        doc.add_paragraph(f"Sentences flagged as AI: {results['flagged_sentences']}")
        doc.add_paragraph(f"Flagged percentage: {stats['flagged_percentage']:.2f}%")
        doc.add_paragraph(f"Average AI probability: {stats['average_ai_probability']*100:.2f}%")
        
        # Color legend
        doc.add_heading('Legend', 2)
        legend = doc.add_paragraph()
        legend.add_run('Green: Low AI probability (<50%)\n').font.color.rgb = RGBColor(0, 128, 0)
        legend.add_run('Yellow: Medium AI probability (50-70%)\n').font.color.rgb = RGBColor(200, 150, 0)
        legend.add_run('Red: High AI probability (>70%)\n').font.color.rgb = RGBColor(200, 0, 0)
        
        # Flagged sentences with humanized versions
        doc.add_page_break()
        doc.add_heading('Flagged Content & Humanized Alternatives', 1)
        
        flagged_count = 0
        for i, sent_data in enumerate(results['sentences_data']):
            if sent_data['needs_humanization']:
                flagged_count += 1
                
                # Section header
                doc.add_heading(f'Section {flagged_count}', 2)
                
                # Original text
                orig_para = doc.add_paragraph()
                orig_para.add_run('ORIGINAL ').bold = True
                orig_run = orig_para.add_run(
                    f"(AI: {sent_data['ai_probability']*100:.1f}%): "
                )
                orig_run.font.color.rgb = RGBColor(200, 0, 0)
                orig_para.add_run(sent_data['original'])
                
                # Humanized versions
                if sent_data['humanized_versions']:
                    doc.add_paragraph().add_run('HUMANIZED OPTIONS:').bold = True
                    
                    for j, version in enumerate(sent_data['humanized_versions'][:3], 1):
                        option_para = doc.add_paragraph()
                        option_para.add_run(f'Option {j} ').bold = True
                        
                        # Color code based on AI probability
                        if version['ai_probability'] < 0.5:
                            color = RGBColor(0, 128, 0)  # Green
                        elif version['ai_probability'] < 0.7:
                            color = RGBColor(200, 150, 0)  # Yellow
                        else:
                            color = RGBColor(200, 0, 0)  # Red
                        
                        ai_run = option_para.add_run(
                            f"(AI: {version['ai_probability']*100:.1f}%, "
                            f"Improvement: {version['improvement']*100:.1f}%): "
                        )
                        ai_run.font.color.rgb = color
                        option_para.add_run(version['text'])
                
                doc.add_paragraph()  # Spacing
        
        # Full humanized document
        doc.add_page_break()
        doc.add_heading('Complete Humanized Document', 1)
        doc.add_paragraph('(Using best humanized version for flagged sentences)')
        doc.add_paragraph()
        
        for sent_data in results['sentences_data']:
            if sent_data['needs_humanization'] and sent_data['humanized_versions']:
                # Use best humanized version (lowest AI probability)
                best_version = sent_data['humanized_versions'][0]['text']
                doc.add_paragraph(best_version)
            else:
                doc.add_paragraph(sent_data['original'])
        
        doc.save(output_file)
        print(f"Document saved: {output_file}")
    
    def create_text_output(self, results: Dict, output_file: str):
        """Create a text file with humanized content"""
        print(f"\nCreating text output: {output_file}")
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("="*80 + "\n")
            f.write("AI CONTENT HUMANIZATION REPORT\n")
            f.write("="*80 + "\n\n")
            
            # Statistics
            stats = results['statistics']
            f.write(f"Total sentences analyzed: {results['total_sentences']}\n")
            f.write(f"Sentences flagged as AI: {results['flagged_sentences']}\n")
            f.write(f"Flagged percentage: {stats['flagged_percentage']:.2f}%\n")
            f.write(f"Average AI probability: {stats['average_ai_probability']*100:.2f}%\n\n")
            
            # Flagged content
            f.write("\n" + "-"*80 + "\n")
            f.write("FLAGGED CONTENT WITH HUMANIZED ALTERNATIVES\n")
            f.write("-"*80 + "\n\n")
            
            flagged_count = 0
            for sent_data in results['sentences_data']:
                if sent_data['needs_humanization']:
                    flagged_count += 1
                    f.write(f"\n{'='*60}\n")
                    f.write(f"SECTION {flagged_count}\n")
                    f.write(f"{'='*60}\n\n")
                    
                    f.write(f"ORIGINAL (AI: {sent_data['ai_probability']*100:.1f}%):\n")
                    f.write(f"{sent_data['original']}\n\n")
                    
                    if sent_data['humanized_versions']:
                        f.write("HUMANIZED OPTIONS:\n\n")
                        for j, version in enumerate(sent_data['humanized_versions'][:3], 1):
                            f.write(f"Option {j} (AI: {version['ai_probability']*100:.1f}%, "
                                   f"Improvement: {version['improvement']*100:.1f}%):\n")
                            f.write(f"{version['text']}\n\n")
            
            # Full humanized text
            f.write("\n" + "="*80 + "\n")
            f.write("COMPLETE HUMANIZED DOCUMENT\n")
            f.write("="*80 + "\n\n")
            
            for sent_data in results['sentences_data']:
                if sent_data['needs_humanization'] and sent_data['humanized_versions']:
                    best_version = sent_data['humanized_versions'][0]['text']
                    f.write(best_version + " ")
                else:
                    f.write(sent_data['original'] + " ")
        
        print(f"Text report saved: {output_file}")
    
    def process_and_humanize(self, pdf_path: str, threshold: float = 0.7,
                            output_docx: str = "humanized_document.docx",
                            output_txt: str = "humanization_report.txt"):
        """Main processing function"""
        # Extract text
        text = self.extract_text_from_pdf(pdf_path)
        
        if not text:
            print("No text extracted from PDF!")
            return
        
        # Process document
        results = self.process_document(text, threshold)
        
        # Display summary
        print("\n" + "="*80)
        print("PROCESSING COMPLETE")
        print("="*80)
        print(f"Total sentences: {results['total_sentences']}")
        print(f"Flagged as AI: {results['flagged_sentences']}")
        print(f"Flagged percentage: {results['statistics']['flagged_percentage']:.2f}%")
        print(f"Average AI probability: {results['statistics']['average_ai_probability']*100:.2f}%")
        
        # Create outputs
        self.create_comparison_docx(results, output_docx)
        self.create_text_output(results, output_txt)
        
        print("\n✓ Humanization complete!")
        print(f"✓ Review {output_docx} for side-by-side comparison")
        print(f"✓ Review {output_txt} for text-only report")


def main():
    parser = argparse.ArgumentParser(
        description='Detect AI content and humanize it using paraphrasing'
    )
    parser.add_argument('pdf_path', type=str, 
                       help='Path to the PDF file to process')
    parser.add_argument('--threshold', type=float, default=0.7,
                       help='AI detection threshold (0-1, default: 0.7)')
    parser.add_argument('--output-docx', type=str, default='humanized_document.docx',
                       help='Output Word document name')
    parser.add_argument('--output-txt', type=str, default='humanization_report.txt',
                       help='Output text report name')
    parser.add_argument('--detector', type=str, default='roberta-base-openai-detector',
                       help='AI detection model')
    parser.add_argument('--paraphraser', type=str, 
                       default='humarin/chatgpt_paraphraser_on_T5_base',
                       help='Paraphrasing model')
    
    args = parser.parse_args()
    
    humanizer = AIContentHumanizer(
        detector_model=args.detector,
        paraphrase_model=args.paraphraser
    )
    
    humanizer.process_and_humanize(
        args.pdf_path,
        threshold=args.threshold,
        output_docx=args.output_docx,
        output_txt=args.output_txt
    )


if __name__ == "__main__":
    main()
